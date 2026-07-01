"""Smoke test: Word (.docx) tracked changes via _docx_revisions.

Self-contained (no Word needed): it builds a fixture per revision type — the authored tier
(insert_tracked / delete_tracked) plus an injected-fragment tier for every other type, using the
exact markup real Word emits — then asserts BOTH accept and reject for each, that no revision is
left afterwards, that only word/document.xml changed, and that python-docx can still open the file.

The accept/reject expectations here are the ones a live Microsoft Word round-trip produced during
development (see dev/make_fixtures.py); this test guards them without needing Word installed.

Run: python tests/smoke_docx_revisions.py
"""
import os
import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))
import _docx_revisions as R          # noqa: E402
from _ooxml_zip import patch_parts   # noqa: E402
from docx import Document            # noqa: E402
from lxml import etree               # noqa: E402

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
DOC = "word/document.xml"
_tmp = tempfile.mkdtemp(prefix="docx_rev_")
_n = [0]


def _w(t):
    return f"{{{W}}}{t}"


def _m(t):
    return f"{{{M}}}{t}"


def _rev(el, i=0):
    el.set(_w("id"), str(i))
    el.set(_w("author"), "Reviewer")
    el.set(_w("date"), "2026-01-01T00:00:00Z")
    return el


def _run(text, deleted=False, math=False):
    r = etree.Element(_m("r") if math else _w("r"))
    t = etree.SubElement(r, _m("t") if math else _w(("delText" if deleted else "t")))
    t.text = text
    if not math:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _wrap(tag, *kids):
    el = _rev(etree.Element(_w(tag)))
    for k in kids:
        el.append(k)
    return el


def _newpath():
    _n[0] += 1
    return os.path.join(_tmp, f"t{_n[0]}.docx")


def _para_doc(path, n=1):
    d = Document()
    for _ in range(n):
        d.add_paragraph()
    d.save(path)


def _first_p(root):
    return root.find(f".//{_w('p')}")


def _all_p(root):
    body = root.find(_w("body"))
    return [c for c in body if c.tag == _w("p")]


def _set_children(p, *els):
    for r in p.findall(_w("r")):
        p.remove(r)
    for e in els:
        p.append(e)


def _inject(path, fn):
    def mut(ps):
        root = ps.get_xml(DOC)
        fn(root)
        ps.set_xml(DOC, root)
    patch_parts(path, mut)


# ---- probes -------------------------------------------------------------
def _parts(path):
    with zipfile.ZipFile(path) as z:
        return {i.filename: z.read(i.filename) for i in z.infolist()}


def txt(path):
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read(DOC))
    return "".join(t.text or "" for t in root.iter(_w("t")))


def mtxt(path):
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read(DOC))
    return "".join(t.text or "" for t in root.iter(_w("t"), _m("t")))


def npara(path):
    with zipfile.ZipFile(path) as z:
        return len(_all_p(etree.fromstring(z.read(DOC))))


def nrows(path):
    with zipfile.ZipFile(path) as z:
        tbl = etree.fromstring(z.read(DOC)).find(f".//{_w('tbl')}")
    return len(tbl.findall(_w("tr")))


def ncells1(path):
    with zipfile.ZipFile(path) as z:
        tbl = etree.fromstring(z.read(DOC)).find(f".//{_w('tbl')}")
    return len(tbl.findall(_w("tr"))[0].findall(_w("tc")))


def fmt(path):
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read(DOC))
    return "bold" if root.find(f".//{_w('b')}") is not None else "plain"


def jc(path):
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read(DOC))
    j = root.find(f".//{_w('jc')}")
    return j.get(_w("val")) if j is not None else "left"


def opens(path):
    try:
        Document(path)
        return True
    except Exception as e:  # noqa: BLE001
        return f"{type(e).__name__}: {e}"


# ---- fixture builders ---------------------------------------------------
def build_content_ins(path):
    _para_doc(path)
    _inject(path, lambda root: _set_children(_first_p(root), _run("Keep"), _wrap("ins", _run("New"))))


def build_content_del(path):
    _para_doc(path)
    _inject(path, lambda root: _set_children(_first_p(root), _run("Keep"),
                                             _wrap("del", _run("Gone", deleted=True))))


def build_rprchange(path):
    _para_doc(path)

    def fn(root):
        r = etree.Element(_w("r"))
        rpr = etree.SubElement(r, _w("rPr"))
        etree.SubElement(rpr, _w("b"))
        chg = _rev(etree.SubElement(rpr, _w("rPrChange")))
        etree.SubElement(chg, _w("rPr"))   # old = not bold
        t = etree.SubElement(r, _w("t"))
        t.text = "Word"
        _set_children(_first_p(root), r)
    _inject(path, fn)


def build_pprchange(path):
    _para_doc(path)

    def fn(root):
        p = _first_p(root)
        ppr = etree.SubElement(p, _w("pPr"))
        p.insert(0, ppr)
        etree.SubElement(ppr, _w("jc")).set(_w("val"), "center")
        chg = _rev(etree.SubElement(ppr, _w("pPrChange")))
        etree.SubElement(chg, _w("pPr"))   # old = default (left)
        p.append(_run("Body"))
    _inject(path, fn)


def _paramark(path, tag):
    _para_doc(path, 2)

    def fn(root):
        p1, p2 = _all_p(root)
        ppr = etree.SubElement(p1, _w("pPr"))
        p1.insert(0, ppr)
        rpr = etree.SubElement(ppr, _w("rPr"))
        _rev(etree.SubElement(rpr, _w(tag)))
        p1.append(_run("First"))
        p2.append(_run("Second"))
    _inject(path, fn)


def _row_mark(path, tag):
    d = Document()
    t = d.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            t.cell(r, c).text = f"r{r}c{c}"
    d.save(path)

    def fn(root):
        tr = root.find(f".//{_w('tbl')}").findall(_w("tr"))[1]
        trpr = tr.find(_w("trPr"))
        if trpr is None:
            trpr = etree.Element(_w("trPr"))
            tr.insert(0, trpr)
        _rev(etree.SubElement(trpr, _w(tag)))
    _inject(path, fn)


def build_move(path):
    _para_doc(path)

    def fn(root):
        p = _first_p(root)
        for r in p.findall(_w("r")):
            p.remove(r)
        fs = _rev(etree.SubElement(p, _w("moveFromRangeStart")))
        fs.set(_w("name"), "m1")
        p.append(_wrap("moveFrom", _run("MOVED")))
        etree.SubElement(p, _w("moveFromRangeEnd")).set(_w("id"), "0")
        p.append(_run("REST"))
        ts = _rev(etree.SubElement(p, _w("moveToRangeStart")), 1)
        ts.set(_w("name"), "m1")
        p.append(_wrap("moveTo", _run("MOVED")))
        etree.SubElement(p, _w("moveToRangeEnd")).set(_w("id"), "1")
    _inject(path, fn)


def build_math_del(path):
    _para_doc(path)

    def fn(root):
        p = _first_p(root)
        for r in p.findall(_w("r")):
            p.remove(r)
        om = etree.SubElement(p, _m("oMath"))
        r1 = etree.SubElement(om, _m("r"))
        r1.append(_rev(etree.Element(_w("del"))))       # deleted math run marker
        t1 = etree.SubElement(r1, _m("t"))
        t1.text = "a"
        r2 = etree.SubElement(om, _m("r"))
        t2 = etree.SubElement(r2, _m("t"))
        t2.text = "+b"
    _inject(path, fn)


def _cell_mark(path, tag):
    d = Document()
    t = d.add_table(rows=1, cols=3)
    for c in range(3):
        t.cell(0, c).text = f"c{c}"
    d.save(path)

    def fn(root):
        tc = root.find(f".//{_w('tbl')}").findall(_w("tr"))[0].findall(_w("tc"))[2]
        tcpr = tc.find(_w("tcPr"))
        _rev(etree.SubElement(tcpr, _w(tag)))
    _inject(path, fn)


def build_authored_ins(path):
    d = Document()
    d.add_paragraph("Hello world")
    d.save(path)
    R.insert_tracked(path, {"text": "Hello"}, " dear", author="Reviewer")


def build_authored_del(path):
    d = Document()
    d.add_paragraph("Hello cruel world")
    d.save(path)
    R.delete_tracked(path, {"text": "cruel "}, author="Reviewer")


# ---- scenario runner ----------------------------------------------------
def check(name, build, probe, exp_accept, exp_reject, fails):
    for action, exp in (("accept", exp_accept), ("reject", exp_reject)):
        p = _newpath()
        build(p)
        before = _parts(p)
        (R.accept_all if action == "accept" else R.reject_all)(p)
        after = _parts(p)
        changed = {n for n in set(before) | set(after) if before.get(n) != after.get(n)}
        if changed - {DOC}:
            fails.append(f"{name}/{action}: touched {sorted(changed - {DOC})}")
        left = R.list_revisions(p)
        if left:
            fails.append(f"{name}/{action}: {len(left)} revision(s) remain")
        o = opens(p)
        if o is not True:
            fails.append(f"{name}/{action}: python-docx cannot open ({o})")
        got = probe(p)
        if got != exp:
            fails.append(f"{name}/{action}: got {got!r}, expected {exp!r}")


def main() -> int:
    fails = []
    check("content-insert", build_content_ins, txt, "KeepNew", "Keep", fails)
    check("content-delete", build_content_del, txt, "Keep", "KeepGone", fails)
    check("run-format", build_rprchange, fmt, "bold", "plain", fails)
    check("para-format", build_pprchange, jc, "center", "left", fails)
    check("inserted-para-mark", lambda p: _paramark(p, "ins"), npara, 2, 1, fails)
    check("deleted-para-mark", lambda p: _paramark(p, "del"), npara, 1, 2, fails)
    check("row-insert", lambda p: _row_mark(p, "ins"), nrows, 2, 1, fails)
    check("row-delete", lambda p: _row_mark(p, "del"), nrows, 1, 2, fails)
    check("move", build_move, txt, "RESTMOVED", "MOVEDREST", fails)
    check("math-delete", build_math_del, mtxt, "+b", "a+b", fails)
    check("cell-insert", lambda p: _cell_mark(p, "cellIns"), ncells1, 3, 2, fails)
    check("cell-delete", lambda p: _cell_mark(p, "cellDel"), ncells1, 2, 3, fails)
    check("authored-insert", build_authored_ins, txt, "Hello dear world", "Hello world", fails)
    check("authored-delete", build_authored_del, txt, "Hello world", "Hello cruel world", fails)

    # toggle: on adds <w:trackRevisions/>, off removes it (only settings.xml changes)
    tp = _newpath()
    _para_doc(tp)
    before = _parts(tp)
    R.set_tracking(tp, True)
    with zipfile.ZipFile(tp) as z:
        on = etree.fromstring(z.read("word/settings.xml")).find(_w("trackRevisions")) is not None
    if not on:
        fails.append("set_tracking(on): trackRevisions not present")
    if {n for n in _parts(tp) if _parts(tp)[n] != before.get(n)} - {"word/settings.xml"}:
        fails.append("set_tracking(on): touched parts other than settings.xml")
    R.set_tracking(tp, False)
    with zipfile.ZipFile(tp) as z:
        off = etree.fromstring(z.read("word/settings.xml")).find(_w("trackRevisions")) is None
    if not off:
        fails.append("set_tracking(off): trackRevisions still present")

    if fails:
        print("FAIL:")
        for f in fails:
            print("  -", f)
        return 1
    print("PASS - docx tracked changes (list / accept / reject / toggle / author, every type)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
