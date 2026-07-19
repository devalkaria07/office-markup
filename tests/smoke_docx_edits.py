"""Smoke test: v0.3.0 Word editing commands — replace_tracked, direct edits, apply_edits.

Asserts: tracked replace = deletion + insertion with consecutive ids (accept -> new text,
reject -> original); direct edits leave NO revisions and never touch settings.xml; a
direct edit in a header changes only that part; the batch applies sequentially (a later
anchor can match text an earlier edit inserted), is all-or-nothing on failure, validates
the whole spec upfront, and --dry-run leaves the file byte-identical.

Run: python tests/smoke_docx_edits.py
"""
import os
import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))
import _docx_batch as B              # noqa: E402
import _docx_comments as dc          # noqa: E402
import _docx_edit as E               # noqa: E402
import _docx_revisions as R          # noqa: E402
import _errors                       # noqa: E402
from _ooxml_zip import patch_parts   # noqa: E402
from docx import Document            # noqa: E402
from lxml import etree               # noqa: E402

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DOC = "word/document.xml"
HDR = "word/header1.xml"
_tmp = tempfile.mkdtemp(prefix="docx_edits_")
_n = [0]


def _w(t):
    return f"{{{W}}}{t}"


def _newpath():
    _n[0] += 1
    return os.path.join(_tmp, f"t{_n[0]}.docx")


def _base(path, text="Hello cruel world."):
    d = Document()
    d.add_paragraph(text)
    d.save(path)


def _parts(path):
    with zipfile.ZipFile(path) as z:
        return {i.filename: z.read(i.filename) for i in z.infolist()}


def _text(path, part=DOC):
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read(part))
    return "".join(t.text or "" for t in root.iter(_w("t")))


def _opens(path):
    try:
        Document(path)
        return True
    except Exception as e:  # noqa: BLE001
        return f"{type(e).__name__}: {e}"


def _add_header(path, text):
    def mut(ps):
        hdr = etree.Element(_w("hdr"), nsmap={"w": W})
        p = etree.SubElement(hdr, _w("p"))
        r = etree.SubElement(p, _w("r"))
        t = etree.SubElement(r, _w("t"))
        t.text = text
        ps.add_xml_part(HDR, hdr)
        ps.ensure_content_type(
            "/" + HDR, "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml")
        rid = ps.rels_for(DOC).add_rel(RNS + "/header", "header1.xml")
        root = ps.get_xml(DOC)
        sect = root.find(_w("body")).find(_w("sectPr"))
        ref = etree.Element(_w("headerReference"))
        ref.set(_w("type"), "default")
        ref.set(f"{{{RNS}}}id", rid)
        sect.insert(0, ref)
        ps.set_xml(DOC, root)
    patch_parts(path, mut)


def main() -> int:
    fails = []

    # --- replace_tracked: one deletion + one insertion, consecutive ids ---------------
    f = _newpath()
    _base(f)
    R.replace_tracked(f, {"text": "cruel"}, "kind", author="Reviewer")
    recs = R.list_revisions(f)
    kinds = sorted(r["type"] for r in recs)
    if kinds != ["deletion", "insertion"]:
        fails.append(f"replace: expected deletion+insertion, got {kinds}")
    with zipfile.ZipFile(f) as z:
        root = etree.fromstring(z.read(DOC))
    ids = [int(el.get(_w("id"))) for tag in ("del", "ins") for el in root.iter(_w(tag))]
    if len(ids) != 2 or abs(ids[0] - ids[1]) != 1:
        fails.append(f"replace: revision ids not consecutive: {ids}")
    R.accept_all(f)
    if _text(f) != "Hello kind world.":
        fails.append(f"replace/accept: got {_text(f)!r}")
    f = _newpath()
    _base(f)
    R.replace_tracked(f, {"text": "cruel"}, "kind", author="Reviewer")
    R.reject_all(f)
    if _text(f) != "Hello cruel world.":
        fails.append(f"replace/reject: got {_text(f)!r}")

    # --- direct edits: text changes, ZERO revisions, settings.xml untouched -----------
    cases = (
        ("replace", lambda p: E.replace_direct(p, {"text": "cruel"}, "kind"), "Hello kind world."),
        ("insert-after", lambda p: E.insert_direct(p, {"text": "world"}, "!"), "Hello cruel world!."),
        ("insert-before", lambda p: E.insert_direct(p, {"text": "Hello"}, ">> ", before=True),
         ">> Hello cruel world."),
        ("delete", lambda p: E.delete_direct(p, {"text": "cruel "}), "Hello world."),
    )
    for name, fn, expected in cases:
        f = _newpath()
        _base(f)
        before = _parts(f)
        fn(f)
        after = _parts(f)
        if _text(f) != expected:
            fails.append(f"direct {name}: got {_text(f)!r}, expected {expected!r}")
        if R.list_revisions(f):
            fails.append(f"direct {name}: created revisions")
        changed = {n for n in set(before) | set(after) if before.get(n) != after.get(n)}
        if changed != {DOC}:
            fails.append(f"direct {name}: changed {sorted(changed)}, expected only document.xml")
        if (o := _opens(f)) is not True:
            fails.append(f"direct {name}: won't open: {o}")

    # --- direct edit in a header story touches only that part -------------------------
    f = _newpath()
    _base(f)
    _add_header(f, "Company Confidential")
    before = _parts(f)
    E.replace_direct(f, {"text": "Confidential"}, "Public")
    after = _parts(f)
    changed = {n for n in set(before) | set(after) if before.get(n) != after.get(n)}
    if changed != {HDR}:
        fails.append(f"direct header edit: changed {sorted(changed)}, expected only {HDR}")
    if _text(f, HDR) != "Company Public":
        fails.append(f"direct header edit: header text {_text(f, HDR)!r}")

    # --- apply_edits: sequential semantics + mixed actions ----------------------------
    f = _newpath()
    _base(f, "The Q3 figures need review.")
    cid = dc.add_comment(f, {"text": "review"}, "Existing thread.", author="Alex Morgan")
    spec = {
        "author": "Alex Morgan",
        "edits": [
            # direct edits (tracked: false) change the anchorable text immediately...
            {"action": "replace", "anchor": {"text": "Q3 figures"}, "text": "Q3 FINAL figures",
             "tracked": False},
            # ...so this anchor only exists AFTER edit 1 ran -> proves sequential application
            # (text inside a TRACKED insertion is deliberately NOT anchorable — it is
            # pending, not real, until accepted)
            {"action": "insert", "anchor": {"text": "FINAL"}, "text": " (approved)",
             "tracked": False},
            {"action": "delete", "anchor": {"text": "need "}},          # tracked by default
            {"action": "comment", "anchor": {"text": "figures"}, "text": "Who signs off?"},
            {"action": "reply", "comment_id": cid, "text": "Done in this pass."},
            {"action": "resolve", "comment_id": cid},
        ],
    }
    out = B.apply_edits(f, spec)
    if out["applied"] != 6:
        fails.append(f"batch: applied {out['applied']}, expected 6")
    new_ids = [r["id"] for r in out["results"] if "id" in r]
    if len(new_ids) != 2:
        fails.append(f"batch: expected 2 created ids (comment+reply), got {new_ids}")
    kinds = [r["type"] for r in R.list_revisions(f)]
    if kinds != ["deletion"]:
        fails.append(f"batch: expected exactly one tracked deletion pending, got {kinds}")
    R.accept_all(f)
    if "Q3 FINAL (approved) figures review." not in _text(f):
        fails.append(f"batch: text after accept: {_text(f)!r}")
    recs = {r["id"]: r for r in dc.list_comments(f)}
    if cid not in recs or not recs[cid]["resolved"]:
        fails.append("batch: original thread not resolved")
    if not any(r["parent_id"] == cid for r in recs.values()):
        fails.append("batch: reply not threaded under the original comment")
    if (o := _opens(f)) is not True:
        fails.append(f"batch: won't open: {o}")

    # --- all-or-nothing: a failing middle edit leaves the file byte-identical ---------
    f = _newpath()
    _base(f)
    before = _parts(f)
    bad = {"author": "A", "edits": [
        {"action": "insert", "anchor": {"text": "Hello"}, "text": " there"},
        {"action": "delete", "anchor": {"text": "no-such-phrase-xyz"}},
        {"action": "insert", "anchor": {"text": "world"}, "text": "!"},
    ]}
    try:
        B.apply_edits(f, bad)
        fails.append("all-or-nothing: expected AnchorNotFound")
    except _errors.AnchorNotFound as e:
        if "edit 2" not in str(e):
            fails.append(f"all-or-nothing: error does not name the failing edit: {e}")
    except Exception as e:  # noqa: BLE001
        fails.append(f"all-or-nothing: wrong error {type(e).__name__}: {e}")
    if _parts(f) != before:
        fails.append("all-or-nothing: file was modified despite the failure")

    # --- validate: EVERY spec problem reported at once, nothing touched ---------------
    f = _newpath()
    _base(f)
    before = _parts(f)
    broken = {"edits": [
        {"action": "frobnicate"},
        {"action": "replace", "anchor": {"text": "Hello"}, "text": "Hi"},   # no author anywhere
        {"action": "reply", "text": "hi"},                                  # no comment_id, no author
    ]}
    try:
        B.apply_edits(f, broken)
        fails.append("validate: expected CommentError")
    except _errors.CommentError as e:
        msg = str(e)
        for frag in ("edit 1", "unknown action", "edit 2", "author", "edit 3", "comment_id"):
            if frag not in msg:
                fails.append(f"validate: aggregated message missing {frag!r}")
    if _parts(f) != before:
        fails.append("validate: file was modified by an invalid spec")

    # --- review fixes: spans crossing a container REFUSE loudly (never silent skips) ---
    def _linked_doc(path):
        d = Document()
        d.add_paragraph()
        d.save(path)

        def mut(ps):
            root = ps.get_xml(DOC)
            p = root.find(f".//{_w('p')}")
            for text in ("Click ",):
                r = etree.SubElement(p, _w("r"))
                t = etree.SubElement(r, _w("t"))
                t.text = text
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            link = etree.SubElement(p, _w("hyperlink"))
            link.set(_w("anchor"), "top")
            lr = etree.SubElement(link, _w("r"))
            lt = etree.SubElement(lr, _w("t"))
            lt.text = "this link"
            r2 = etree.SubElement(p, _w("r"))
            t2 = etree.SubElement(r2, _w("t"))
            t2.text = " now"
            t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            ps.set_xml(DOC, root)
        patch_parts(path, mut)

    for name, call in (
        ("tracked delete", lambda p: R.delete_tracked(p, {"text": "Click this link now"},
                                                      author="A")),
        ("tracked replace", lambda p: R.replace_tracked(p, {"text": "Click this link now"},
                                                        "NEW", author="A")),
        ("direct delete", lambda p: E.delete_direct(p, {"text": "Click this link now"})),
        ("direct replace", lambda p: E.replace_direct(p, {"text": "Click this link now"}, "NEW")),
    ):
        f = _newpath()
        _linked_doc(f)
        before = _parts(f)
        try:
            call(f)
            fails.append(f"mid-span link / {name}: expected AnchorNotFound, call succeeded")
        except _errors.AnchorNotFound as e:
            if "separately" not in str(e):
                fails.append(f"mid-span link / {name}: message lacks guidance: {e}")
        except Exception as e:  # noqa: BLE001
            fails.append(f"mid-span link / {name}: wrong error {type(e).__name__}: {e}")
        if _parts(f) != before:
            fails.append(f"mid-span link / {name}: file modified despite refusal")

    # paragraph anchor on a mixed text+link paragraph: paragraph-specific guidance
    f = _newpath()
    _linked_doc(f)
    try:
        R.delete_tracked(f, {"paragraph": 1}, author="A")
        fails.append("mixed-para anchor: expected AnchorNotFound")
    except _errors.AnchorNotFound as e:
        if "by phrase" not in str(e):
            fails.append(f"mixed-para anchor: message lacks paragraph guidance: {e}")

    # batch: "part" on a comment action is rejected upfront (comments are body-anchored)
    f = _newpath()
    _base(f)
    try:
        B.apply_edits(f, {"author": "A", "edits": [
            {"action": "comment", "anchor": {"text": "Hello", "part": "header1"}, "text": "x"}]})
        fails.append("batch comment part: expected CommentError")
    except _errors.CommentError as e:
        if "body-anchored" not in str(e):
            fails.append(f"batch comment part: message wrong: {e}")

    # --- dry run: full execution in memory, file byte-identical -----------------------
    f = _newpath()
    _base(f)
    before = _parts(f)
    out = B.apply_edits(f, {"author": "A", "edits": [
        {"action": "replace", "anchor": {"text": "cruel"}, "text": "kind"}]}, dry_run=True)
    if out["applied"] != 1 or not out["dry_run"]:
        fails.append(f"dry-run: bad result {out}")
    if _parts(f) != before:
        fails.append("dry-run: file was modified")

    if fails:
        print("FAIL:")
        for x in fails:
            print("  -", x)
        return 1
    print("PASS - docx edits (tracked replace / direct edits / batch apply_edits)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
