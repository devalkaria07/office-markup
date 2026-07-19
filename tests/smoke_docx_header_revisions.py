"""Smoke test: tracked changes in header / footnote stories (#8) + the settings.xml
bootstrap on authoring (#4) + cross-story authoring (v0.3.0).

Self-contained (no Word needed): a header part and a footnotes part are injected with the
exact markup Word emits, then list / accept / reject / author are asserted across stories.

Run: python tests/smoke_docx_header_revisions.py
"""
import os
import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))
import _docx_revisions as R          # noqa: E402
import _errors                       # noqa: E402
from _ooxml_zip import patch_parts   # noqa: E402
from docx import Document            # noqa: E402
from lxml import etree               # noqa: E402

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DOC = "word/document.xml"
HDR = "word/header1.xml"
FNS = "word/footnotes.xml"
_tmp = tempfile.mkdtemp(prefix="docx_hdr_rev_")
_n = [0]


def _w(t):
    return f"{{{W}}}{t}"


def _newpath():
    _n[0] += 1
    return os.path.join(_tmp, f"t{_n[0]}.docx")


def _run(text):
    r = etree.Element(_w("r"))
    t = etree.SubElement(r, _w("t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _ins(rev_id, *kids):
    el = etree.Element(_w("ins"))
    el.set(_w("id"), str(rev_id))
    el.set(_w("author"), "Reviewer")
    el.set(_w("date"), "2026-01-01T00:00:00Z")
    for k in kids:
        el.append(k)
    return el


def _para(*kids):
    p = etree.Element(_w("p"))
    for k in kids:
        p.append(k)
    return p


def _build(path, *, body_ins=False, header=False, header_ins=False, footnote_ins=False,
           header_text="Header line"):
    """A doc with optional revisions in body / header / footnotes, built by injection."""
    d = Document()
    d.add_paragraph("Body text here.")
    d.save(path)

    def mut(ps):
        root = ps.get_xml(DOC)
        body = root.find(_w("body"))
        if body_ins:
            p = body.find(_w("p"))
            p.append(_ins(1, _run(" plus body change")))
        if header or header_ins:
            hdr = etree.Element(_w("hdr"), nsmap={"w": W})
            hp = _para(_run(header_text))
            if header_ins:
                hp.append(_ins(2, _run(" plus header change")))
            hdr.append(hp)
            ps.add_xml_part(HDR, hdr)
            ps.ensure_content_type(
                "/" + HDR,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml")
            rid = ps.rels_for(DOC).add_rel(RNS + "/header", "header1.xml")
            sect = body.find(_w("sectPr"))
            ref = etree.Element(_w("headerReference"))
            ref.set(_w("type"), "default")
            ref.set(f"{{{RNS}}}id", rid)
            sect.insert(0, ref)
        if footnote_ins:
            fns = etree.Element(_w("footnotes"), nsmap={"w": W})
            fn = etree.SubElement(fns, _w("footnote"))
            fn.set(_w("id"), "2")
            fp = _para(_run("A footnote."), _ins(3, _run(" plus note change")))
            fn.append(fp)
            ps.add_xml_part(FNS, fns)
            ps.ensure_content_type(
                "/" + FNS,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
            ps.rels_for(DOC).add_rel(RNS + "/footnotes", "footnotes.xml")
            p = body.find(_w("p"))
            fr = etree.Element(_w("r"))
            fref = etree.SubElement(fr, _w("footnoteReference"))
            fref.set(_w("id"), "2")
            p.append(fr)
        ps.set_xml(DOC, root)
    patch_parts(path, mut)


def _parts(path):
    with zipfile.ZipFile(path) as z:
        return {i.filename: z.read(i.filename) for i in z.infolist()}


def _part_text(path, part):
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read(part))
    return "".join(t.text or "" for t in root.iter(_w("t")))


def _has_ins(path, part):
    with zipfile.ZipFile(path) as z:
        return b"<w:ins " in z.read(part)


def _opens(path):
    try:
        Document(path)
        return True
    except Exception as e:  # noqa: BLE001
        return f"{type(e).__name__}: {e}"


def main() -> int:
    fails = []

    # --- 1) list sees every story, in fixed order, with part/story fields -------------
    f = _newpath()
    _build(f, body_ins=True, header_ins=True, footnote_ins=True)
    recs = R.list_revisions(f)
    if len(recs) != 3:
        fails.append(f"list: expected 3 revisions across stories, got {len(recs)}")
    else:
        got = [(r["id"], r["part"], r["story"]) for r in recs]
        want = [(1, DOC, "document"), (2, HDR, "header"), (3, FNS, "footnotes")]
        if got != want:
            fails.append(f"list: bad order/parts: {got}")
        if "header1.xml" not in recs[1]["location"]:
            fails.append(f"list: header location missing part name: {recs[1]['location']!r}")

    # --- 2) accept ONE header revision by global id: only header1.xml changes ---------
    before = _parts(f)
    R.accept(f, 2)
    after = _parts(f)
    changed = {n for n in set(before) | set(after) if before.get(n) != after.get(n)}
    if changed != {HDR}:
        fails.append(f"accept(header): changed {sorted(changed)}, expected only {HDR}")
    if _part_text(f, HDR) != "Header line plus header change":
        fails.append(f"accept(header): header text {_part_text(f, HDR)!r}")
    if len(R.list_revisions(f)) != 2:
        fails.append("accept(header): expected 2 revisions to remain")

    # --- 3) accept_all / reject_all sweep every story (the #8 regression) -------------
    f = _newpath()
    _build(f, body_ins=True, header_ins=True, footnote_ins=True)
    n = R.accept_all(f)
    if n != 3:
        fails.append(f"accept_all: reported {n}, expected 3")
    if R.list_revisions(f):
        fails.append("accept_all: revisions remain")
    for part in (HDR, FNS):
        if _has_ins(f, part):
            fails.append(f"accept_all: w:ins left behind in {part}")
    if (o := _opens(f)) is not True:
        fails.append(f"accept_all: won't open: {o}")
    f = _newpath()
    _build(f, body_ins=True, header_ins=True, footnote_ins=True)
    R.reject_all(f)
    if R.list_revisions(f):
        fails.append("reject_all: revisions remain")
    if _part_text(f, HDR) != "Header line":
        fails.append(f"reject_all: header text {_part_text(f, HDR)!r}")

    # --- 4) authoring INTO a header (cross-story anchor), only that part + settings ---
    f = _newpath()
    _build(f, header=True, header_text="Confidential draft")
    before = _parts(f)
    R.insert_tracked(f, {"text": "Confidential draft"}, " v2", author="Reviewer")
    after = _parts(f)
    changed = {n for n in set(before) | set(after) if before.get(n) != after.get(n)}
    if changed - {HDR, "word/settings.xml"}:
        fails.append(f"author-in-header: changed {sorted(changed - {HDR, 'word/settings.xml'})}")
    recs = R.list_revisions(f)
    if len(recs) != 1 or recs[0]["part"] != HDR:
        fails.append(f"author-in-header: bad listing {[(r['id'], r['part']) for r in recs]}")
    R.accept_all(f)
    if _part_text(f, HDR) != "Confidential draft v2":
        fails.append(f"author-in-header: header text {_part_text(f, HDR)!r}")

    # --- 5) cross-story ambiguity names the parts; a part pin resolves it -------------
    f = _newpath()
    _build(f, header=True, header_text="Body text here.")   # same phrase in body + header
    try:
        R.delete_tracked(f, {"text": "Body text here."}, author="Reviewer")
        fails.append("ambiguity: expected AmbiguousAnchor")
    except _errors.AmbiguousAnchor as e:
        msg = str(e)
        if "document.xml" not in msg or "header1.xml" not in msg:
            fails.append(f"ambiguity: message lacks part names: {msg!r}")
    except Exception as e:  # noqa: BLE001
        fails.append(f"ambiguity: wrong error {type(e).__name__}: {e}")
    R.delete_tracked(f, {"text": "Body text here.", "part": "header1"}, author="Reviewer")
    R.accept_all(f)
    if _part_text(f, HDR) != "":
        fails.append(f"part-pin: header text {_part_text(f, HDR)!r}, expected ''")
    if "Body text here." not in _part_text(f, DOC):
        fails.append("part-pin: body text was touched")

    # --- 6) #4: authoring bootstraps settings.xml when the package has none -----------
    for fn_name in ("insert_tracked", "delete_tracked"):
        f = _newpath()
        d = Document()
        d.add_paragraph("Some anchor text.")
        d.save(f)

        def strip_settings(ps):
            ps.drop_part("word/settings.xml")
            ps.rels_for(DOC).remove_by_target("settings.xml")
        patch_parts(f, strip_settings)
        fn = getattr(R, fn_name)
        if fn_name == "insert_tracked":
            fn(f, {"text": "anchor"}, " more", author="Reviewer")
        else:
            fn(f, {"text": "anchor "}, author="Reviewer")
        with zipfile.ZipFile(f) as z:
            names = set(z.namelist())
            if "word/settings.xml" not in names:
                fails.append(f"{fn_name}: settings.xml not created")
                continue
            sroot = etree.fromstring(z.read("word/settings.xml"))
            if sroot.find(_w("trackRevisions")) is None:
                fails.append(f"{fn_name}: trackRevisions missing from created settings.xml")
            ct = z.read("[Content_Types].xml").decode()
            if "wordprocessingml.settings+xml" not in ct:
                fails.append(f"{fn_name}: settings content type not registered")
            rels = z.read("word/_rels/document.xml.rels").decode()
            if "relationships/settings" not in rels:
                fails.append(f"{fn_name}: settings relationship not wired")
        if (o := _opens(f)) is not True:
            fails.append(f"{fn_name}: won't open: {o}")

    if fails:
        print("FAIL:")
        for x in fails:
            print("  -", x)
        return 1
    print("PASS - docx header/footnote revisions (#8) + settings bootstrap (#4) + cross-story authoring")
    return 0


if __name__ == "__main__":
    sys.exit(main())
