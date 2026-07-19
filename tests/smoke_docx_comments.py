"""Smoke test: modern threaded comments in Word (.docx).

Builds a tiny doc, then drives add -> reply -> resolve/reopen -> delete entirely through
_docx_comments, asserting structure, threading, the resolved flag, anchor round-trip, the
byte-stability guarantee (only comment-related parts change), and that python-docx still
opens every intermediate file.

Run: python tests/smoke_docx_comments.py
"""
import os
import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))
import _docx_comments as dc          # noqa: E402
import _errors                       # noqa: E402
from _ooxml_zip import patch_parts   # noqa: E402
from docx import Document            # noqa: E402
from lxml import etree               # noqa: E402

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(t):
    return f"{{{W}}}{t}"


def _mk_run(text):
    r = etree.Element(_w("r"))
    t = etree.SubElement(r, _w("t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _inject(path, fn):
    def mut(ps):
        root = ps.get_xml("word/document.xml")
        fn(root)
        ps.set_xml("word/document.xml", root)
    patch_parts(path, mut)


def _base(path):
    d = Document()
    d.add_paragraph("The Q3 revenue figures need review before we circulate this report.")
    d.add_paragraph("A second paragraph mentions revenue again to test ambiguity.")
    d.add_heading("Scope", level=1)
    d.save(path)


def _opens(path):
    try:
        Document(path)
        return True
    except Exception as e:
        return f"{type(e).__name__}: {e}"


def _parts(path):
    with zipfile.ZipFile(path) as z:
        return {i.filename: z.read(i.filename) for i in z.infolist()}


def _docxml(path):
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8", "ignore")


def main() -> int:
    fails = []
    tmp = tempfile.mkdtemp(prefix="docx_comments_")
    f = os.path.join(tmp, "t.docx")
    _base(f)
    before = _parts(f)

    # --- ambiguity + not-found guards (before we touch anything) ---
    try:
        dc.add_comment(f, {"text": "revenue"}, "x", author="A")
        fails.append("ambiguous anchor: expected AmbiguousAnchor")
    except _errors.AmbiguousAnchor:
        pass
    except Exception as e:
        fails.append(f"ambiguous anchor: wrong error {type(e).__name__}: {e}")
    try:
        dc.add_comment(f, {"text": "does-not-exist-xyz"}, "x", author="A")
        fails.append("missing anchor: expected AnchorNotFound")
    except _errors.AnchorNotFound:
        pass
    except Exception as e:
        fails.append(f"missing anchor: wrong error {type(e).__name__}: {e}")

    # --- add ---
    cid = dc.add_comment(f, {"text": "revenue", "occurrence": 1},
                         "Please confirm this figure matches the latest forecast.",
                         author="Alex Morgan")
    recs = dc.list_comments(f)
    if len(recs) != 1:
        fails.append(f"add: expected 1 comment, got {len(recs)}")
    else:
        r = recs[0]
        if r["anchor_text"] != "revenue":
            fails.append(f"add: anchor_text={r['anchor_text']!r}")
        if r["author"] != "Alex Morgan" or r["initials"] != "AM":
            fails.append(f"add: author/initials wrong: {r['author']!r}/{r['initials']!r}")
        if r["resolved"] or r["is_reply"]:
            fails.append("add: should be unresolved top-level")
        if "Q3 revenue figures" not in r["context"]:
            fails.append(f"add: context missing surrounding text: {r['context']!r}")
    with zipfile.ZipFile(f) as z:
        names = set(z.namelist())
        for p in ("word/comments.xml", "word/commentsExtended.xml", "word/commentsIds.xml",
                  "word/commentsExtensible.xml", "word/people.xml"):
            if p not in names:
                fails.append(f"add: missing part {p}")
        ct = z.read("[Content_Types].xml").decode()
        if "wordprocessingml.comments+xml" not in ct:
            fails.append("add: comments content-type not registered")
        rels = z.read("word/_rels/document.xml.rels").decode()
        for frag in ("relationships/comments", "office/2011/relationships/commentsExtended",
                     "office/2016/09/relationships/commentsIds",
                     "office/2018/08/relationships/commentsExtensible",
                     "office/2011/relationships/people"):
            if frag not in rels:
                fails.append(f"add: relationship missing: {frag}")
    if (o := _opens(f)) is not True:
        fails.append(f"add: won't open: {o}")
    after = _parts(f)
    changed = set(n for n in before if before[n] != after.get(n))
    # settings.xml changes because we upgrade compatibilityMode 14 -> 15 so Word allows Resolve
    allowed = {"word/document.xml", "[Content_Types].xml", "word/_rels/document.xml.rels", "word/settings.xml"}
    if not changed.issubset(allowed):
        fails.append(f"add: changed unexpected existing parts: {changed - allowed}")
    if "word/document.xml" not in changed or "word/settings.xml" not in changed:
        fails.append(f"add: expected document.xml + settings.xml to change, got {sorted(changed)}")

    # --- reply ---
    rid = dc.reply(f, cid, "Confirmed - it matches the final forecast.", author="Sam Lee")
    recs = {r["id"]: r for r in dc.list_comments(f)}
    if len(recs) != 2:
        fails.append(f"reply: expected 2 comments, got {len(recs)}")
    if rid in recs:
        rep = recs[rid]
        if rep["parent_id"] != cid or not rep["is_reply"] or rep["thread_id"] != cid:
            fails.append(f"reply: bad threading parent={rep['parent_id']} thread={rep['thread_id']}")
        if rep["author"] != "Sam Lee":
            fails.append("reply: author wrong")
    else:
        fails.append("reply: reply id not listed")
    if _docxml(f).count("commentReference") != 2:
        fails.append("reply: expected 2 commentReference anchors")
    if (o := _opens(f)) is not True:
        fails.append(f"reply: won't open: {o}")

    # --- resolve / reopen (via either id in the thread) ---
    dc.set_status(f, cid, True)
    if not all(r["resolved"] for r in dc.list_comments(f)):
        fails.append("resolve: thread not marked resolved")
    dc.set_status(f, rid, False)   # reopen via the reply id -> same thread
    if any(r["resolved"] for r in dc.list_comments(f)):
        fails.append("reopen: thread still resolved")
    if (o := _opens(f)) is not True:
        fails.append(f"resolve: won't open: {o}")

    # --- delete the reply, then the root ---
    dc.delete(f, rid)
    recs = dc.list_comments(f)
    if [r["id"] for r in recs] != [cid]:
        fails.append(f"delete reply: expected [{cid}], got {[r['id'] for r in recs]}")
    if _docxml(f).count("commentReference") != 1:
        fails.append("delete reply: expected 1 remaining anchor")
    dc.delete(f, cid)
    if dc.list_comments(f):
        fails.append("delete root: expected 0 comments")
    dx = _docxml(f)
    for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
        if tag in dx:
            fails.append(f"delete root: {tag} left behind in document.xml")
    if (o := _opens(f)) is not True:
        fails.append(f"delete: won't open: {o}")

    # --- v0.3.0: comment anchored on hyperlink text (#5) ---
    f2 = os.path.join(tmp, "h.docx")
    d = Document()
    d.add_paragraph()
    d.save(f2)

    def link_para(root):
        p = root.find(f".//{_w('p')}")
        link = etree.Element(_w("hyperlink"))
        link.set(_w("anchor"), "top")
        link.append(_mk_run("this link"))
        for el in (_mk_run("Click "), link, _mk_run(" now")):
            p.append(el)
    _inject(f2, link_para)
    try:
        dc.add_comment(f2, {"text": "this link"}, "Check the link target.", author="Alex Morgan")
    except Exception as e:  # noqa: BLE001
        fails.append(f"hyperlink comment: raised {type(e).__name__}: {e}")
    else:
        rl = dc.list_comments(f2)
        if len(rl) != 1 or rl[0]["anchor_text"] != "this link":
            fails.append(f"hyperlink comment: bad listing {[(r['id'], r['anchor_text']) for r in rl]}")
        if (o := _opens(f2)) is not True:
            fails.append(f"hyperlink comment: won't open: {o}")

    # --- v0.3.0: comment in a multi-w:t run leaves the break outside the range (#6) ---
    f3 = os.path.join(tmp, "m.docx")
    d = Document()
    d.add_paragraph()
    d.save(f3)

    def multi_t_para(root):
        p = root.find(f".//{_w('p')}")
        r = _mk_run("Hello ")
        etree.SubElement(r, _w("br"))
        t2 = etree.SubElement(r, _w("t"))
        t2.text = "world"
        p.append(r)
    _inject(f3, multi_t_para)
    try:
        dc.add_comment(f3, {"text": "Hello"}, "Greeting check.", author="Alex Morgan")
    except Exception as e:  # noqa: BLE001
        fails.append(f"multi-t comment: raised {type(e).__name__}: {e}")
    else:
        with zipfile.ZipFile(f3) as z:
            root = etree.fromstring(z.read("word/document.xml"))
        in_range, br_inside = False, False
        for el in root.iter():
            if el.tag == _w("commentRangeStart"):
                in_range = True
            elif el.tag == _w("commentRangeEnd"):
                in_range = False
            elif el.tag == _w("br") and in_range:
                br_inside = True
        if br_inside:
            fails.append("multi-t comment: w:br ended up inside the comment range")
        rl = dc.list_comments(f3)
        if len(rl) != 1 or rl[0]["anchor_text"] != "Hello":
            fails.append(f"multi-t comment: bad listing {[(r['id'], r['anchor_text']) for r in rl]}")
        if (o := _opens(f3)) is not True:
            fails.append(f"multi-t comment: won't open: {o}")

    # --- v0.3.0: classic (pre-2013, no w14:paraId) comments and delete (#7) ---
    def _inject_classic(path, cid="9"):
        def mut(ps):
            if ps.has("word/comments.xml"):
                comments = ps.get_xml("word/comments.xml")
            else:
                comments = etree.Element(_w("comments"), nsmap={"w": W})
                ps.add_xml_part("word/comments.xml", comments)
                ps.ensure_content_type(
                    "/word/comments.xml",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
                ps.rels_for("word/document.xml").add_rel(
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
                    "comments.xml")
            cm = etree.SubElement(comments, _w("comment"))
            cm.set(_w("id"), cid)
            cm.set(_w("author"), "Old Reviewer")
            cm.set(_w("initials"), "OR")
            cp = etree.SubElement(cm, _w("p"))          # deliberately NO w14:paraId
            cr = etree.SubElement(cp, _w("r"))
            ct_ = etree.SubElement(cr, _w("t"))
            ct_.text = "legacy note"
            ps.set_xml("word/comments.xml", comments)
            root = ps.get_xml("word/document.xml")
            para = root.find(f".//{_w('p')}")
            st = etree.Element(_w("commentRangeStart"))
            st.set(_w("id"), cid)
            en = etree.Element(_w("commentRangeEnd"))
            en.set(_w("id"), cid)
            para.insert(0, st)
            para.append(en)
            rr = etree.SubElement(para, _w("r"))
            cref = etree.SubElement(rr, _w("commentReference"))
            cref.set(_w("id"), cid)
            ps.set_xml("word/document.xml", root)
        patch_parts(path, mut)

    def _mixed_fixture(path):
        _base(path)
        root_id = dc.add_comment(path, {"text": "Q3 revenue figures"},
                                 "Modern thread root.", author="Alex Morgan")
        rep_id = dc.reply(path, root_id, "Modern reply.", author="Sam Lee")
        _inject_classic(path)
        return root_id, rep_id

    # (a) deleting the CLASSIC comment removes it and ONLY it (old code nuked the
    #     modern threads instead and left the classic in place)
    f4 = os.path.join(tmp, "c1.docx")
    root_id, rep_id = _mixed_fixture(f4)
    dc.delete(f4, "9")
    left = {r["id"] for r in dc.list_comments(f4)}
    if left != {root_id, rep_id}:
        fails.append(f"classic delete (mixed): expected modern thread to survive, got {sorted(left)}")
    if "commentReference" in _docxml(f4) and _docxml(f4).count("commentReference") != 2:
        fails.append("classic delete (mixed): anchor count wrong after delete")
    if (o := _opens(f4)) is not True:
        fails.append(f"classic delete (mixed): won't open: {o}")

    # (b) deleting the MODERN root leaves the classic comment untouched
    f5 = os.path.join(tmp, "c2.docx")
    root_id, rep_id = _mixed_fixture(f5)
    dc.delete(f5, root_id)
    left = [r["id"] for r in dc.list_comments(f5)]
    if left != ["9"]:
        fails.append(f"modern delete (mixed): expected only classic '9' left, got {left}")

    # (c) classic-only file WITHOUT commentsExtended: delete works (old: silent no-op)
    f6 = os.path.join(tmp, "c3.docx")
    _base(f6)
    _inject_classic(f6)
    dc.delete(f6, "9")
    if dc.list_comments(f6):
        fails.append("classic-only delete: comment still listed (silent no-op)")
    dx6 = _docxml(f6)
    for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
        if tag in dx6:
            fails.append(f"classic-only delete: {tag} left behind")
    if (o := _opens(f6)) is not True:
        fails.append(f"classic-only delete: won't open: {o}")

    # (d) resolve / reply on a classic comment raise a CLEAR CommentError
    f7 = os.path.join(tmp, "c4.docx")
    _base(f7)
    _inject_classic(f7)
    for act, call in (("resolve", lambda: dc.set_status(f7, "9", True)),
                      ("reply", lambda: dc.reply(f7, "9", "hi", author="A"))):
        try:
            call()
            fails.append(f"classic {act}: expected CommentError")
        except _errors.CommentNotFound as e:
            fails.append(f"classic {act}: misleading CommentNotFound: {e}")
        except _errors.CommentError as e:
            if "classic" not in str(e):
                fails.append(f"classic {act}: message lacks 'classic': {e}")
        except Exception as e:  # noqa: BLE001
            fails.append(f"classic {act}: wrong error {type(e).__name__}: {e}")

    # --- v0.3.0: atomic writes leave no temp residue (U0) ---
    residue = [n for n in os.listdir(tmp) if n.endswith(".om-tmp")]
    if residue:
        fails.append(f"atomic write: temp residue left behind: {residue}")

    if fails:
        print("FAIL:")
        for x in fails:
            print("  -", x)
        return 1
    print("PASS — docx threaded comments (add / reply / resolve / reopen / delete)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
